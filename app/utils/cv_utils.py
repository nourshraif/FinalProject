"""
CV text extraction for common resume formats.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Optional, Set

from app.utils.pdf_utils import extract_text_from_pdf

ALLOWED_CV_EXTENSIONS: Set[str] = {".pdf", ".docx", ".doc", ".txt", ".rtf"}
ALLOWED_CV_LABEL = "PDF, Word (.doc/.docx), RTF, or plain text (.txt)"


def cv_extension(filename: Optional[str]) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def is_allowed_cv_filename(filename: Optional[str]) -> bool:
    ext = cv_extension(filename)
    return f".{ext}" in ALLOWED_CV_EXTENSIONS if ext else False


def _extract_text_from_docx(content: bytes) -> Optional[str]:
    try:
        from docx import Document

        doc = Document(BytesIO(content))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        text = "\n".join(parts).strip()
        return text or None
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


def _extract_text_from_txt(content: bytes) -> Optional[str]:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = content.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    return None


def _extract_text_from_rtf(content: bytes) -> Optional[str]:
    try:
        from striprtf.striprtf import rtf_to_text

        raw = content.decode("latin-1", errors="ignore")
        text = rtf_to_text(raw).strip()
        return text or None
    except Exception as e:
        print(f"Error extracting text from RTF: {e}")
        return None


def _extract_text_from_doc(content: bytes) -> Optional[str]:
    """Best-effort extraction for legacy Word .doc (OLE) files."""
    # Mislabeled .docx files sometimes use a .doc extension.
    docx_text = _extract_text_from_docx(content)
    if docx_text:
        return docx_text

    try:
        import olefile

        if not olefile.isOleFile(BytesIO(content)):
            return None
        ole = olefile.OleFileIO(BytesIO(content))
        if not ole.exists("WordDocument"):
            ole.close()
            return None
        stream = ole.openstream("WordDocument").read()
        ole.close()
        raw = stream.decode("latin-1", errors="ignore")
        chunks = re.findall(r"[\x20-\x7E]{4,}", raw)
        text = "\n".join(chunks).strip()
        return text or None
    except Exception as e:
        print(f"Error extracting text from DOC: {e}")
        return None


def extract_text_from_cv_bytes(content: bytes, filename: str = "") -> Optional[str]:
    """Extract plain text from CV file bytes."""
    ext = cv_extension(filename)
    if ext == "pdf":
        return extract_text_from_pdf(BytesIO(content))
    if ext == "docx":
        return _extract_text_from_docx(content)
    if ext == "doc":
        return _extract_text_from_doc(content)
    if ext == "txt":
        return _extract_text_from_txt(content)
    if ext == "rtf":
        return _extract_text_from_rtf(content)
    return None


def extract_text_from_cv_file(
    file_obj,
    filename: str = "",
) -> Optional[str]:
    """Extract plain text from a file-like CV upload."""
    name = filename or getattr(file_obj, "name", "") or ""
    try:
        pos = file_obj.tell()
    except Exception:
        pos = None
    try:
        data = file_obj.read()
        if isinstance(data, str):
            data = data.encode("utf-8", errors="ignore")
        return extract_text_from_cv_bytes(data, name)
    finally:
        if pos is not None:
            try:
                file_obj.seek(pos)
            except Exception:
                pass
